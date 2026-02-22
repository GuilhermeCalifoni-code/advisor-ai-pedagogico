from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from app.models.schemas import DocumentState

def create_rfc_document(doc_data: DocumentState) -> BytesIO:
    document = Document()

    # Estilo do Título Principal
    title = document.add_heading('REQUISIÇÃO DE MUDANÇA (RfC)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Informações Gerais
    document.add_heading('1. INFORMAÇÕES GERAIS', level=1)
    
    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Preenchendo a tabela
    row0 = table.rows[0].cells
    row0[0].text = f"TÍTULO:\n{doc_data.titulo}"
    row0[1].text = f"ÁREA SOLICITANTE:\n{doc_data.areaSolicitante}"
    
    row1 = table.rows[1].cells
    row1[0].text = f"RESPONSÁVEL:\n{doc_data.responsavel}"
    row1[1].text = f"DATA:\n{doc_data.dataSolicitacao}"
    
    row2 = table.rows[2].cells
    # Mescla a última linha para o Sistema Impactado
    cell_merged = row2[0].merge(row2[1])
    cell_merged.text = f"SISTEMA/PROCESSO IMPACTADO:\n{doc_data.sistemaImpactado}"

    document.add_paragraph() # Espaço

    # 2. Levantamento Inicial
    document.add_heading('2. LEVANTAMENTO INICIAL', level=1)

    # Função auxiliar para criar seções bonitas
    def add_section(title, content):
        h = document.add_heading(title, level=2)
        h.style.font.color.rgb = RGBColor(0, 51, 102) # Azul escuro profissional
        p = document.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_section('2.1 Cenário Atual', doc_data.cenarioAtual)
    add_section('2.2 Necessidade do Negócio', doc_data.necessidadeNegocio)
    add_section('2.3 Benefícios Esperados', doc_data.beneficiosEsperados)
    add_section('2.4 Expectativa Funcional', doc_data.expectativaFuncional)

    # Salva na memória (não no disco, para ser rápido)
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream