import os
import json
from pathlib import Path
from io import BytesIO

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv

# --- LER .ENV ---
env_path = Path(__file__).parent / '.env'
load_dotenv(dotenv_path=env_path, override=True)

try:
    from langchain_core.prompts import PromptTemplate
    from langchain_core.messages import HumanMessage
except ImportError:
    from langchain.prompts import PromptTemplate
    from langchain.schema import HumanMessage

from langchain_google_genai import ChatGoogleGenerativeAI

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- 1. CONFIGURAÇÃO DA IA ---
api_key = os.getenv("GOOGLE_API_KEY")

if not api_key:
    print("❌ ERRO CRÍTICO: Chave GOOGLE_API_KEY não encontrada no arquivo .env!")

llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    temperature=0.4, 
    google_api_key=api_key
)

# --- 2. NOVA ESTRUTURA DO DOCUMENTO ---
class DocumentState(BaseModel):
    professor: str = ""
    data: str = ""
    segmento: str = ""
    competencia_foco: str = ""
    resumo_necessidade: str = ""
    plano_acao_detalhado: str = ""

class ChatRequest(BaseModel):
    message: str
    current_doc: DocumentState
    conversation_history: list = []
    user_name: str = "Coordenador"
    attachment: str | None = None

# --- 3. PROMPT FOCADO E BLINDADO CONTRA ERROS ---
# --- 3. PROMPT FOCADO E BLINDADO CONTRA ERROS ---
chat_prompt = PromptTemplate.from_template("""
Você é um Advisor Pedagógico Sênior especialista em Educação Bilíngue (CLIL).
Seu ÚNICO objetivo é criar um PLANO DE AÇÃO DE DESENVOLVIMENTO PROFISSIONAL detalhado.

ESTADO ATUAL DO DOCUMENTO:
{current_doc}

MENSAGEM:
{input}

---
DIRETRIZES DE GERAÇÃO:
Crie o texto do 'plano_acao_detalhado' OBRIGATORIAMENTE nesta estrutura:

I. Pesquisa e Formação
1. Aprofundamento no tema: (Parágrafo inspirador sobre a competência).
Material para estudo:
- Livro: [Título e Autor] - [Por que esta leitura é fundamental?]
- Vídeo: [Título do vídeo/Canal] - [Qual o insight principal?]
- Podcast: [Episódio ou nome] - [Como isso ajuda?]
- Artigo: [Conceito ou blog] - [Técnica a extrair]

II. Aplicação Prática
2. Probing Questions: (5 a 7 perguntas reflexivas profundas).
3. Ações Imediatas: (3 passos práticos para a próxima aula amanhã).

III. Mentoria e Acompanhamento
4. Indicadores de Sucesso: (O que será observado na próxima visita do Advisor).

IV. Plano de Estudo Personalizado
5. Plano de Estudo Semanal:
Para as próximas 4 semanas. IMPORTANTE: NÃO use formato de tabela com barras (|). Formate como uma lista executiva, limpa e altamente profissional, usando EXATAMENTE este padrão:

SEMANA 1
• Atividade: [Descreva a leitura, vídeo ou prática da semana]
• Propósito: [Explique o objetivo e o impacto esperado]

SEMANA 2
• Atividade: [Descreva a leitura, vídeo ou prática da semana]
• Propósito: [Explique o objetivo e o impacto esperado]

SEMANA 3
• Atividade: [Descreva a leitura, vídeo ou prática da semana]
• Propósito: [Explique o objetivo e o impacto esperado]

SEMANA 4
• Atividade: [Descreva a leitura, vídeo ou prática da semana]
• Propósito: [Explique o objetivo e o impacto esperado]

---
ATENÇÃO - REGRAS TÉCNICAS CRÍTICAS (OBRIGATÓRIO):
1. NÃO USE FORMATAÇÃO MARKDOWN NO TEXTO. Não use asteriscos (**) para negrito, não use hashtags (#). Use APENAS texto puro e simples.
2. Formate a sua resposta como um JSON válido. 
3. Para quebras de linha dentro do texto gerado, use rigorosamente '\\n'. Nunca utilize um "Enter" real no meio da string JSON.

SAÍDA ESPERADA (JSON PURO):
{{
  "ai_message": "Seu feedback rápido e encorajador para o coordenador...",
  "updated_doc": {{ ... TODOS os campos do DocumentState atualizados ... }}
}}
""")

@app.post("/api/chat")
async def chat_endpoint(request: ChatRequest):
    try:
        # Prepara o histórico
        formatted_history = "\n".join([f"{msg['role']}: {msg['content']}" for msg in request.conversation_history[-5:] if isinstance(msg['content'], str)])
        
        # Formata o prompt
        text_content = chat_prompt.format(
            current_doc=request.current_doc.model_dump(), 
            history=formatted_history, 
            input=request.message, 
            user_name=request.user_name
        )
        
        # Prepara o payload multimodal
        message_payload = [{"type": "text", "text": text_content}]
        if request.attachment:
            message_payload.append({"type": "image_url", "image_url": request.attachment})

        # Invoca a IA
        response = llm.invoke([HumanMessage(content=message_payload)])
        
        # Limpa formatação markdown de blocos de código e espaços extras
        content_str = response.content.replace("```json", "").replace("```", "").strip()
        
        # strict=False permite ao Python ler o JSON mesmo com quebras de linha (Enters) acidentais
        return json.loads(content_str, strict=False)
        
    except Exception as e:
        print(f"❌ Erro Crítico no JSON: {e}")
        try:
            print(f"RESPOSTA BRUTA DA IA (para debug): {response.content}")
        except:
            pass
        return {
            "ai_message": "A IA gerou um plano tão complexo que causou um pequeno erro de formatação. Por favor, tente enviar a mensagem novamente.", 
            "updated_doc": request.current_doc
        }

# --- 5. ENDPOINT DE DOWNLOAD DO WORD (.DOCX) ---
from fastapi.responses import StreamingResponse
from docx import Document

def create_action_plan_document(doc_data: DocumentState):
    doc = Document()
    doc.add_heading('Plano de Ação - Desenvolvimento Docente', 0)
    
    doc.add_heading('1. IDENTIFICAÇÃO', level=1)
    doc.add_paragraph(f"Professor(a): {doc_data.professor}\nData: {doc_data.data}\nSegmento: {doc_data.segmento}")
    
    doc.add_heading('2. FOCO DE DESENVOLVIMENTO', level=1)
    doc.add_paragraph(f"Competência Alvo: {doc_data.competencia_foco}\nResumo da Necessidade: {doc_data.resumo_necessidade}")

    doc.add_heading('3. PLANO DE AÇÃO ESTRUTURADO', level=1)
    doc.add_paragraph(doc_data.plano_acao_detalhado)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.post("/api/download")
async def download_doc(doc: DocumentState):
    file_stream = create_action_plan_document(doc)
    filename = f"PlanoAcao_{doc.professor.replace(' ', '_')[:20] or 'Professor'}.docx"
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)